"""
PDF ↔ Word Converter - Public Web App
Dùng ConvertAPI + Supabase để lưu stats + giới hạn 8 lần/ngày
"""

import os, uuid, threading, time, requests
from datetime import date
from flask import Flask, render_template, request, send_file, jsonify
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 20 * 1024 * 1024  # 20MB

UPLOAD_FOLDER = "/tmp/uploads"
OUTPUT_FOLDER = "/tmp/outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

ALLOWED = {".pdf", ".docx", ".doc"}
CONVERTAPI_SECRET = os.environ.get("CONVERTAPI_SECRET", "")
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")
DAILY_LIMIT = 8


def supabase_headers():
    return {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation"
    }


def get_stats():
    try:
        r = requests.get(
            f"{SUPABASE_URL}/rest/v1/stats?id=eq.1",
            headers=supabase_headers(), timeout=5
        )
        data = r.json()
        if data:
            return data[0]
    except:
        pass
    return {"total": 0, "pdf_to_word": 0, "word_to_pdf": 0, "today_count": 0, "last_reset": str(date.today())}


def check_and_increment(kind):
    """Kiểm tra giới hạn ngày và tăng bộ đếm. Trả về (ok, stats, error_msg)"""
    try:
        current = get_stats()
        today = str(date.today())

        # Reset nếu sang ngày mới
        if current.get("last_reset") != today:
            current["today_count"] = 0
            current["last_reset"] = today

        # Kiểm tra giới hạn
        if current["today_count"] >= DAILY_LIMIT:
            return False, current, f"Đã đạt giới hạn {DAILY_LIMIT} lần/ngày. Vui lòng thử lại vào ngày mai!"

        # Tăng bộ đếm
        current["total"] += 1
        current[kind] += 1
        current["today_count"] += 1

        requests.patch(
            f"{SUPABASE_URL}/rest/v1/stats?id=eq.1",
            headers=supabase_headers(),
            json={
                "total": current["total"],
                kind: current[kind],
                "today_count": current["today_count"],
                "last_reset": today
            },
            timeout=5
        )
        return True, current, None
    except Exception as e:
        return True, {}, None  # Nếu lỗi DB thì vẫn cho convert


def auto_delete(path, delay=600):
    def _del():
        time.sleep(delay)
        try: os.remove(path)
        except: pass
    threading.Thread(target=_del, daemon=True).start()


def convert_via_convertapi(input_path, output_path, from_fmt, to_fmt):
    try:
        url = f"https://v2.convertapi.com/convert/{from_fmt}/to/{to_fmt}?Secret={CONVERTAPI_SECRET}"
        with open(input_path, "rb") as f:
            files = {"File": (os.path.basename(input_path), f)}
            r = requests.post(url, files=files, timeout=120)
        if r.status_code != 200:
            return False, f"ConvertAPI lỗi: {r.text}"
        import base64
        data = r.json()
        file_data = data["Files"][0]["FileData"]
        with open(output_path, "wb") as f:
            f.write(base64.b64decode(file_data))
        return True, "ConvertAPI"
    except Exception as e:
        return False, str(e)


def pdf_to_word(pdf_path, docx_path):
    if CONVERTAPI_SECRET:
        return convert_via_convertapi(pdf_path, docx_path, "pdf", "docx")
    try:
        import pdfplumber
        from docx import Document
        doc = Document()
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages):
                if i > 0: doc.add_page_break()
                text = page.extract_text() or ""
                for line in text.split("\n"):
                    doc.add_paragraph(line)
        doc.save(docx_path)
        return True, "pdfplumber"
    except Exception as e:
        return False, str(e)


def word_to_pdf(docx_path, pdf_path):
    if CONVERTAPI_SECRET:
        return convert_via_convertapi(docx_path, pdf_path, "docx", "pdf")
    try:
        from docx import Document as DocxDoc
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import cm
        doc_word = DocxDoc(docx_path)
        styles = getSampleStyleSheet()
        story = []
        for para in doc_word.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 0.3*cm)); continue
            sname = para.style.name
            if "Heading 1" in sname: story.append(Paragraph(text, styles["Heading1"]))
            elif "Heading 2" in sname: story.append(Paragraph(text, styles["Heading2"]))
            else: story.append(Paragraph(text, styles["Normal"]))
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        pdf_doc.build(story)
        return True, "reportlab"
    except Exception as e:
        return False, str(e)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/quota")
def quota():
    try:
        r = requests.get(
            f"https://v2.convertapi.com/user?Secret={CONVERTAPI_SECRET}",
            timeout=5
        )
        data = r.json()
        used = data.get("ConversionsConsumed", 0)
        total = data.get("ConversionsTotal", 250)
        remaining = total - used
        return jsonify({"used": used, "total": total, "remaining": remaining})
    except Exception as e:
        return jsonify({"used": 0, "total": 250, "remaining": 250})


@app.route("/stats")
def stats():
    s = get_stats()
    today = str(date.today())
    if s.get("last_reset") != today:
        s["today_count"] = 0
    s["daily_limit"] = DAILY_LIMIT
    s["remaining"] = max(0, DAILY_LIMIT - s.get("today_count", 0))
    return jsonify(s)


@app.route("/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "Không có file."}), 400
    file = request.files["file"]
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ALLOWED:
        return jsonify({"error": "Chỉ hỗ trợ .pdf hoặc .docx"}), 400

    uid = uuid.uuid4().hex[:8]
    safe = secure_filename(file.filename)
    input_path = os.path.join(UPLOAD_FOLDER, f"{uid}_{safe}")
    file.save(input_path)
    auto_delete(input_path)

    if ext == ".pdf":
        out_name = os.path.splitext(safe)[0] + ".docx"
        out_path = os.path.join(OUTPUT_FOLDER, f"{uid}_{out_name}")
        kind = "pdf_to_word"
        label = "Word (.docx)"
    else:
        out_name = os.path.splitext(safe)[0] + ".pdf"
        out_path = os.path.join(OUTPUT_FOLDER, f"{uid}_{out_name}")
        kind = "word_to_pdf"
        label = "PDF"

    # Kiểm tra giới hạn ngày
    ok_limit, current_stats, err_msg = check_and_increment(kind)
    if not ok_limit:
        return jsonify({"error": err_msg}), 429

    # Convert
    if ext == ".pdf":
        ok, method = pdf_to_word(input_path, out_path)
    else:
        ok, method = word_to_pdf(input_path, out_path)

    if not ok:
        return jsonify({"error": f"Thất bại: {method}"}), 500

    current_stats["daily_limit"] = DAILY_LIMIT
    current_stats["remaining"] = max(0, DAILY_LIMIT - current_stats.get("today_count", 0))

    auto_delete(out_path)
    return jsonify({
        "success": True,
        "download_url": f"/download/{uid}_{out_name}",
        "filename": out_name,
        "method": method,
        "label": label,
        "stats": current_stats,
    })


@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        return "File không còn tồn tại (đã tự xóa sau 10 phút).", 404
    return send_file(path, as_attachment=True, download_name=filename)


@app.route("/preview/<filename>")
def preview(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        return "File không còn tồn tại.", 404
    return send_file(path, as_attachment=False, mimetype="application/pdf")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
